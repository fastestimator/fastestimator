# Copyright 2019 The FastEstimator Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
# ==============================================================================
import inspect
import json
import os
from typing import Callable, List, Union

from docx import Document
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from fastestimator.trace.trace import Trace
from fastestimator.util.data import Data
from fastestimator.util.traceability_util import traceable
from fastestimator.util.util import to_list


@traceable()
class QMSTest(Trace):
    """Automate QMS testing and report generation.

    Args:
        test_descriptions: List of text-based descriptions.
        test_criterias: List of test functions. Function input argument names needs to match keys from the data
            dictionary.
        test_title: Title of the test.
        json_output: Path into which to write the output results JSON.
        doc_output: Path into which to write the output QMS summary report (docx).

    Raises:
        AssertionError: If the number of `test_descriptions` and `test_criteria` do not match.
    """
    def __init__(self,
                 test_descriptions: Union[str, List[str]],
                 test_criterias: Union[List[Callable], Callable],
                 test_title: str = "QMSTest",
                 json_output: str = "",
                 doc_output: str = "") -> None:

        self.json_output = json_output
        self.doc_output = doc_output
        self.test_title = test_title
        self.test_descriptions = to_list(test_descriptions)
        self.test_criterias = to_list(test_criterias)
        assert len(self.test_descriptions) == len(self.test_criterias), "inconsistent input length found"
        all_inputs = set()
        for criteria in self.test_criterias:
            all_inputs.update(inspect.signature(criteria).parameters.keys())
        super().__init__(inputs=all_inputs, mode="test")
        self.total_pass, self.total_fail = 0, 0

    def _initialize_json_summary(self) -> None:
        """Initialize json summary
        """
        self.json_summary = {"title": self.test_title, "stories": []}

    def on_begin(self, data: Data) -> None:
        self._initialize_json_summary()
        self.total_pass, self.total_fail = 0, 0

    def on_epoch_end(self, data: Data) -> None:
        for criteria, description in zip(self.test_criterias, self.test_descriptions):
            story = {"description": description}
            is_passed = criteria(*[data[var_name] for var_name in list(inspect.signature(criteria).parameters.keys())])
            story["passed"] = str(is_passed)

            if is_passed:
                self.total_pass += 1
            else:
                self.total_fail += 1

            self.json_summary["stories"].append(story)

    def on_end(self, data: Data) -> None:
        if self.json_output.endswith(".json"):
            json_path = self.json_output
        else:
            json_path = os.path.join(self.json_output, "QMS.json")
        with open(json_path, 'w') as fp:
            json.dump(self.json_summary, fp, indent=4)
        print("Saved QMS JSON report to {}".format(json_path))

        if self.doc_output.endswith(".docx"):
            doc_path = self.doc_output
        else:
            doc_path = os.path.join(self.doc_output, "QMS_summary.docx")

        doc_summary = _QMSDocx(self.total_pass, self.total_fail)
        doc_summary.save(doc_path)
        print("Saved QMS summary report to {}".format(doc_path))


class _QMSDocx:
    """A class to generate QMS summary report templates given total pass and failure case numbers.

    Args:
        total_pass: Total number of passing QMS tests.
        total_fail: Total number of failing QMS tests.
    """
    def __init__(self, total_pass: int, total_fail: int) -> None:
        self.doc = Document()
        self._write_static_p1()
        self._write_test_result(total_pass, total_fail)
        self._write_static_p2()

    def save(self, output_path: str) -> None:
        """Save document object to disk.

        Args:
            output_path: Saving path.
        """
        self.doc.save(output_path)

    def _write_test_result(self, total_pass: int, total_fail: int) -> None:
        """Write the test result table.

        Args:
            total_pass: Total number of passing QMS tests.
            total_fail: Total number of failing QMS tests.
        """
        total_test = total_pass + total_fail

        table = self.doc.add_table(rows=2, cols=4)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()
                if i == 0:
                    run.bold = True
        table.style = "Table Grid"
        self.fill_table(table,
                        [["Model", "Total Tests", "Tests Passed", "Tests Failed"],
                         ["Model#1", str(total_test), str(total_pass), str(total_fail)]])

    def _write_static_p1(self) -> None:
        """Write the first part of the report before test result table.
        """
        para = self.doc.add_paragraph()
        self.add_line_break(para, 9, font_size=Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(22)
        run.add_text("Verification Summary report for Model")
        self.add_line_break(para, 4, font_size=Pt(14))

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("Record of changes")
        run.add_break()

        table = self.doc.add_table(rows=2, cols=4)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()
                run.bold = True
                run.font.size = Pt(14)

        table.style = "Table Grid"
        self.fill_table(table, [["Rev number", "Date", "Author", "Comments"], ["1", "", "<Name>", "Initial revision"]])

        para = self.doc.add_paragraph()
        self.add_line_break(para, 3, Pt(14))
        run = para.add_run()
        run.bold = True
        run.add_text("NOTE:")
        run = para.add_run()
        run.add_text(" Copies for use are available via the MyWorkshop system. Any printed copies are considered"
                     "uncontrolled. All approval signatures are captured electronically in MyWorkshop.")
        self.add_line_break(para, 9, Pt(14))

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("1   Introdction")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.bold = True
        run.add_text("1.1 Purpose & Scope")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_text(
            "This document contains the results of the verification for model for <name>. Tests were executed in "
            "accordance with the associated Verification Plan (DOC). ")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.bold = True
        run.add_text("1.2 References")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_text("Find below all the relevant documents related to any tools used during verification. ")

        table = self.doc.add_table(rows=4, cols=3)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()
                if i == 0:
                    run.bold = True

        table.style = "Table Grid"
        self.fill_table(
            table,
            [["Location", "Reference", "Document Name"], ["Myworkshop", "<DOC>", "Edison AI Model Verification Plan"], [
                "Myworkshop", "<DOC>", "Model Evaluation Tool Validation"
            ], ["Myworkshop", "<DOC>", "The CRS documents are in Approved state"]])

        para = self.doc.add_paragraph()
        self.add_line_break(para, 2, Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("2   Infrastructure Details")

        table = self.doc.add_table(rows=6, cols=2)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()
                if i == 0:
                    run.bold = True
        table.style = "Table Grid"
        self.fill_table(table,
                        [["", "Details"], ["GPU Architecture", ""], ["OS environment", ""], [
                            "Collection ID of test data set in Edison AI Workbench", ""
                        ], ["Model Artifact ID (s)", ""], ["Location of model test scripts", ""]])

        para = self.doc.add_paragraph()
        self.add_line_break(para, 2, Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("3   Verification Tools")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_text("Tools used for verification are listed in Model Evaluation Tool Validation <DOC>")

        para = self.doc.add_paragraph()
        self.add_line_break(para, 2, Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("4   Results of Verification")

        table = self.doc.add_table(rows=2, cols=3)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()

        table.style = "Table Grid"
        self.fill_table(table,
                        [["Document", "Location", "Comments"],
                         ["<DOC>", "Myworkshop", "Verification Procedure is in Approved state"]])

        para = self.doc.add_paragraph()
        self.add_line_break(para, 1, Pt(14))
        run = para.add_run()
        run.bold = True
        run.add_text("4.1 Functional and Performance Tests")

    def _write_static_p2(self) -> None:
        """Write the second part of the report after test result table.
        """
        para = self.doc.add_paragraph()
        self.add_line_break(para, 2, Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("5   Verification Details")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_text("Below table details the summary of the completion of verification cycle. ")

        table = self.doc.add_table(rows=9, cols=2)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()

                if i == 0 or j == 0:
                    run.bold = True

        table.style = "Table Grid"
        self.fill_table(
            table,
            [["Activity", "Details"],
             [
                 "Test set location in ALM",
                 "URL: http://hc-alm12.health.ge.com/qcbin/start_a.jsp "
                 "Domain: SWPE / Project: HealthCloud \n ALM\Test Lab\<location of ALM test set>"
             ], ["Verification Cycle Start Date", ""], ["Verification Cycle End Date", ""],
             ["Name of the Tester(s)", ""], ["Total # of test cases executed", ""], ["Total # of Defects Filed", ""],
             ["Total # of Tests Passed", ""], ["Total # of Tests Failed", ""]])

        para = self.doc.add_paragraph()
        self.add_line_break(para, 2, Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("6   Defect Summary List")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_text("Below table summarizes the defects found during verification cycle."
                     "The defects are tracked in ALM: http://hc-alm12.health.ge.com/qcbin/start_a.jsp"
                     "Domain: SWPE / Project: HealthCloud")

        table = self.doc.add_table(rows=2, cols=5)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()
                if i == 0:
                    run.bold = True

        table.style = "Table Grid"
        self.fill_table(table,
                        [["Defect ID", "Summary", "Classification", "Status", "Justification"], ["", "", "", "", ""]])

        para = self.doc.add_paragraph()
        self.add_line_break(para, 2, Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("7   Verification Deviations")

        table = self.doc.add_table(rows=2, cols=1)
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                run = table.rows[i].cells[j].paragraphs[0].add_run()

        table.style = "Table Grid"
        self.fill_table(table,
                        [["There were no deviations from the verification plan."],
                         ["There were deviations from the verification plan as follows"]])

        para = self.doc.add_paragraph()
        self.add_line_break(para, 2, Pt(14))
        run = para.add_run()
        run.bold = True
        run.font.size = Pt(14)
        run.add_text("8   Conclusion")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_text("The acceptance criteria identified in the Verification plan have been met. All activities "
                     "supporting this verification activity are complete.")

    @staticmethod
    def fill_table(table: Table, content: List[List[str]]) -> None:
        """Fill input `table` object with given `content`.

        Args:
            table: 2-D table object to be filled.
            content: 2-D content to fill the table.

        Raises:
            AssertionError: If the table and content shapes are inconsistent.
        """
        assert len(table.rows) == len(content)
        assert len(table.columns) == len(content[0])

        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                table.rows[i].cells[j].paragraphs[0].runs[0].add_text(content[i][j])

    @staticmethod
    def add_line_break(paragraph: Paragraph, num: int, font_size=None) -> None:
        """Add a number of line breaks into the target `paragraph` object.

        Args:
            paragraph: Target paragraph.
            num: Number of line breaks.
            font_size: Font size of the line break.
        """
        run = paragraph.add_run()
        if font_size:
            run.font.size = font_size

        for i in range(num):
            run.add_break()
